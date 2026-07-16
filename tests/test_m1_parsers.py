# tests/test_m1_parsers.py
import pytest
from pathlib import Path

# pandas 是可选依赖（ExcelParser 需要），未安装时跳过 Excel 相关测试
pytest.importorskip("pandas")

from server.kb.parser.markdown_parser import MarkdownParser
from server.kb.parser.pdf_parser import PdfParser
from server.kb.parser.office_parser import ExcelParser, WordParser
from server.kb.parser.image_parser import ImageParser


def test_markdown_parser_parse():
    parser = MarkdownParser()
    result = parser.parse("# 标题\n\n正文内容")
    assert result.success is True
    assert "标题" in result.content
    assert result.confidence >= 0.85

def test_markdown_parser_plain_text():
    parser = MarkdownParser()
    result = parser.parse("纯文本无格式")
    assert result.success is True
    assert "纯文本无格式" in result.content

def test_excel_parser_basic(tmp_path):
    pytest.importorskip("openpyxl")
    import pandas as pd
    excel_path = tmp_path / "test.xlsx"
    df = pd.DataFrame({"A": [1, 2], "B": ["x", "y"]})
    df.to_excel(excel_path, sheet_name="Sheet1", index=False)

    parser = ExcelParser()
    result = parser.parse(excel_path)
    assert result.success is True
    assert "Sheet1" in result.content
    assert "A" in result.content and "B" in result.content

def test_word_parser_basic(tmp_path):
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx 未安装")

    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("标题", level=1)
    doc.add_paragraph("段落内容")
    doc.save(docx_path)

    parser = WordParser()
    result = parser.parse(docx_path)
    assert result.success is True
    assert "标题" in result.content

def test_image_parser_degradation():
    parser = ImageParser()
    result = parser.parse("fake_image.png")
    assert result.success is True
    assert result.confidence <= 0.3
    assert "待多模态解析" in result.content or "image" in result.content.lower()

def test_pdf_parser_graceful_degradation(tmp_path):
    fake_pdf = tmp_path / "fake.pdf"
    fake_pdf.write_bytes(b"not a real pdf")

    parser = PdfParser()
    result = parser.parse(fake_pdf)
    assert result.success is False or result.confidence < 0.5
